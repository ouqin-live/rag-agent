"""Document loaders for local files and URLs."""

from __future__ import annotations

import hashlib
import logging
import re
from abc import ABC
from pathlib import Path
from urllib.parse import urlparse

from rag_agent.knowledge.base import BaseLoader, Document

logger = logging.getLogger(__name__)


def _make_doc_id(source: str, content: str) -> str:
    """Generate a deterministic document id."""
    return hashlib.sha256(f"{source}:{content}".encode("utf-8")).hexdigest()[:32]


def _parse_markdown_frontmatter(text: str) -> tuple[str, dict[str, str]]:
    """Parse YAML frontmatter from markdown.

    Returns (cleaned_content, metadata_dict). If no frontmatter is present,
    returns (original_text, {}).
    """
    if not text.startswith("---"):
        return text, {}

    parts = text.split("---", 2)
    if len(parts) < 3:
        return text, {}

    frontmatter_raw = parts[1].strip()
    body = parts[2].strip()

    if not frontmatter_raw:
        return body, {}

    try:
        import yaml
        meta = yaml.safe_load(frontmatter_raw)
        if isinstance(meta, dict):
            # 扁平化：list 值转逗号分隔字符串，方便检索过滤
            flat: dict[str, str] = {}
            for k, v in meta.items():
                if isinstance(v, list):
                    flat[str(k)] = ", ".join(str(x) for x in v)
                elif isinstance(v, (str, int, float, bool)):
                    flat[str(k)] = str(v)
            return body, flat
    except Exception as exc:
        logger.debug("Failed to parse YAML frontmatter: %s", exc)

    return body, {}


class TextLoader(BaseLoader, ABC):
    """Load plain text files (.txt, .md)."""

    extensions: tuple[str, ...] = (".txt",)

    def load(self, source: str) -> list[Document]:
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"Text file not found: {source}")
        content = path.read_text(encoding="utf-8")
        doc_id = _make_doc_id(source, content)
        return [
            Document(
                id=doc_id,
                content=content,
                source=source,
                metadata={"type": "text", "filename": path.name},
            )
        ]


class MarkdownLoader(TextLoader):
    """Load markdown files, parse YAML frontmatter into metadata."""

    extensions = (".md", ".markdown")

    def load(self, source: str) -> list[Document]:
        docs = super().load(source)
        for doc in docs:
            doc.content, frontmatter = _parse_markdown_frontmatter(doc.content)
            doc.metadata["type"] = "markdown"
            if frontmatter:
                doc.metadata.update(frontmatter)
        return docs


class PdfLoader(BaseLoader, ABC):
    """Load PDF files using PyMuPDF (fitz).

    Extracts full text including tables when available.
    """

    extensions = (".pdf",)

    def load(self, source: str) -> list[Document]:
        try:
            import fitz  # noqa: F401
        except ImportError as exc:
            raise ImportError(
                "PDF loading requires 'pymupdf'. Install it with: uv add pymupdf"
            ) from exc

        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"PDF file not found: {source}")

        doc = fitz.open(source)
        parts = []
        for page in doc:
            # 1. 文本提取
            text = page.get_text()
            if text:
                parts.append(text)
            # 2. 表格提取（PyMuPDF 原生支持）
            try:
                tables = page.find_tables()
                if tables:
                    for table in tables:
                        rows = table.extract()
                        if rows:
                            table_text = self._format_table(rows)
                            parts.append(table_text)
            except Exception:
                pass  # 表格提取失败不影响整体

        content = "\n".join(parts)
        doc_id = _make_doc_id(source, content)

        # 3. 提取文档元信息
        metadata = {
            "type": "pdf",
            "filename": path.name,
            "pages": len(doc),
        }
        try:
            doc_meta = doc.metadata
            if doc_meta:
                for key in ("title", "author", "subject"):
                    val = doc_meta.get(key)
                    if val:
                        metadata[key] = val
        except Exception:
            pass

        return [
            Document(
                id=doc_id,
                content=content,
                source=source,
                metadata=metadata,
            )
        ]

    @staticmethod
    def _format_table(rows: list[list[str]]) -> str:
        """Format extracted table rows as markdown-like text."""
        if not rows:
            return ""
        lines = [""]
        lines.append(" | ".join(str(c) for c in rows[0]))
        lines.append(" | ".join("---" for _ in rows[0]))
        for row in rows[1:]:
            lines.append(" | ".join(str(c) for c in row))
        lines.append("")
        return "\n".join(lines)


class DocxLoader(BaseLoader, ABC):
    """Load Word (.docx) files using python-docx.

    Extracts paragraphs with heading-level awareness and tables.
    """

    extensions = (".docx",)

    def load(self, source: str) -> list[Document]:
        try:
            from docx import Document as DocxDocument  # noqa: F401
        except ImportError as exc:
            raise ImportError(
                "Word loading requires 'python-docx'. Install it with: uv add python-docx"
            ) from exc

        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"Word file not found: {source}")

        doc = DocxDocument(source)
        parts: list[str] = []

        # 1. 提取段落，保留标题层级
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if "Heading" in style_name or "heading" in style_name:
                # 从样式名中提取标题级数（如 "Heading 1" → #）
                level = 1
                for ch in style_name:
                    if ch.isdigit():
                        level = int(ch)
                        break
                parts.append(f"{'#' * min(level, 3)} {text}")
            else:
                parts.append(text)

        # 2. 提取表格
        for i, table in enumerate(doc.tables):
            parts.append("")
            parts.append(f"【表格 {i + 1}】")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        content = "\n".join(parts)
        doc_id = _make_doc_id(source, content)
        return [
            Document(
                id=doc_id,
                content=content,
                source=source,
                metadata={"type": "docx", "filename": path.name},
            )
        ]


class UrlLoader(BaseLoader, ABC):
    """Load a web page and extract readable text.

    Uses trafilatura for high-quality extraction when available,
    falling back to basic regex-based extraction.
    """

    extensions = ()

    def __init__(self, timeout: int = 15):
        self.timeout = timeout

    def load(self, source: str) -> list[Document]:
        try:
            import requests
        except ImportError as exc:
            raise ImportError(
                "URL loading requires 'requests'. Install it with: uv add requests"
            ) from exc

        parsed = urlparse(source)
        if not parsed.scheme or not parsed.netloc:
            raise ValueError(f"Invalid URL: {source}")

        response = requests.get(source, timeout=self.timeout)
        response.raise_for_status()
        content_type = response.headers.get("Content-Type", "")

        # PDF URL：下载后走 PdfLoader
        if "application/pdf" in content_type:
            return self._load_pdf_url(source, response)

        # 优先 trafilatura，失败降级正则
        text = self._extract_with_trafilatura(response) or self._extract_with_regex(response.text)
        doc_id = _make_doc_id(source, text)
        return [
            Document(
                id=doc_id,
                content=text,
                source=source,
                metadata={"type": "url", "content_type": content_type},
            )
        ]

    @staticmethod
    def _extract_with_trafilatura(response) -> str | None:
        """Use trafilatura for high-quality web text extraction."""
        try:
            import trafilatura

            text = trafilatura.extract(
                response.text,
                url=response.url,
                include_comments=False,
                include_tables=True,
                include_images=False,
                favor_precision=True,
            )
            if text and len(text.strip()) > 20:
                return text.strip()
            return None
        except Exception:
            return None

    @staticmethod
    def _extract_with_regex(html: str) -> str:
        """Fallback: basic HTML-to-text extraction using regex."""
        html = re.sub(r"<script[^>]*>.*?</script>", " ", html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r"<style[^>]*>.*?</style>", " ", html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r"</(p|div|h[1-6]|li|tr|br)\s*>", "\n", html, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"\s+", " ", text).strip()
        return text

    def _load_pdf_url(self, source: str, response) -> list[Document]:
        """Download PDF from URL and delegate to PdfLoader."""
        import tempfile

        parsed = urlparse(source)
        filename = Path(parsed.path).name or "downloaded.pdf"
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(response.content)
            tmp_path = Path(tmp.name)

        try:
            pdf_loader = PdfLoader()
            docs = pdf_loader.load(str(tmp_path))
            for doc in docs:
                doc.source = source
                doc.metadata["filename"] = filename
                doc.metadata["type"] = "url_pdf"
            return docs
        finally:
            try:
                tmp_path.unlink()
            except Exception:
                pass


class HtmlLoader(BaseLoader, ABC):
    """Load local .html / .htm files.

    Reuses UrlLoader's extraction logic for local HTML files.
    """

    extensions = (".html", ".htm")

    def load(self, source: str) -> list[Document]:
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"HTML file not found: {source}")

        html = path.read_text(encoding="utf-8")
        # 优先 trafilatura，失败降级正则
        text = None
        try:
            import trafilatura
            text = trafilatura.extract(
                html,
                include_comments=False,
                include_tables=True,
                include_images=False,
                favor_precision=True,
            )
        except Exception:
            pass

        if not text or len(text.strip()) <= 20:
            text = UrlLoader._extract_with_regex(html)

        doc_id = _make_doc_id(source, text)
        return [
            Document(
                id=doc_id,
                content=text,
                source=source,
                metadata={"type": "html", "filename": path.name},
            )
        ]


class AutoLoader(BaseLoader, ABC):
    """Pick an appropriate loader based on the source string."""

    def __init__(self, timeout: int = 15):
        self.timeout = timeout
        self._loaders = [
            MarkdownLoader(),
            HtmlLoader(),
            DocxLoader(),
            TextLoader(),
            PdfLoader(),
        ]
        self._url_loader = UrlLoader(timeout=timeout)

    def load(self, source: str) -> list[Document]:
        parsed = urlparse(source)
        if parsed.scheme in ("http", "https"):
            return self._url_loader.load(source)

        lower = source.lower()
        for loader in self._loaders:
            if any(lower.endswith(ext) for ext in loader.extensions):
                return loader.load(source)

        # Default to text loader for unknown local paths
        return TextLoader().load(source)
